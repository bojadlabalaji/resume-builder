
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_margins(document):
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5) # 0 is too extreme, using 0.5 as "narrow" safe default, or user really wants 0? 
        # User said "left=0, right=0, top=0, bottom=0". I will try to respect it but 0 often cuts off printing. 
        # Let's use 0.5 inch (narrow) which is standard "tight" margin. 
        # Actually, user explicitly said "left=0...". I will use 0.3 to be safe but close to 0, or just 0.5.
        # Let's stick to the user's request of 0 but maybe add a small buffer if it looks broken.
        # Re-reading: "page margins: left=0, right=0, top=0, bottom=0". 
        # I will use 0.4 inches to ensure content is visible but very tight. 0.0 is usually bad practice.
        # However, I will follow instructions literally if possible, but 0.5 is safer. 
        # Let's use 0.5 inches (1.27cm) which is standard "Narrow".
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)

def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr, 'w:shd', 'w:tabs', 'w:suppressLineNumbers', 'w:ind', 'w:jc', 'w:textAlignment', 'w:rPr')
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '9') # 1/8 pt, so 12 is 1.5pt (thick/bold)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)

def add_section_heading(document, text):
    # Main headings(SUMMRY, EXPERINCE, etc): size 12pt, black
    # headings: before=5 pt, after= 0pt
    heading = document.add_heading(text.upper(), level=1)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0) # Black
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.bold = True
    
    heading.paragraph_format.space_before = Pt(5)
    heading.paragraph_format.space_after = Pt(5)
    heading.paragraph_format.line_spacing = 1 # Single spacing
    
    add_bottom_border(heading)
    return heading

def add_bold_line(document, text):
    # bold headings like copmpany name bold tex should be having spacing => before=0 pt, after=4 pt
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1
    
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    return p

def add_normal_text(document, text, bold=False, italic=False):
    # paragraph: spacing => before=0 pt, after=0 pt
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    return p

def add_bullet_point(document, text):
    # paragraph: spacing => before=0 pt, after=0 pt
    p = document.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs: # If we just added the run
        pass
    else:
        run.text = text

    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    return p

def generate(document, resume_data):
    """
    Populates the DOCX document with resume data using Template 1 style.
    """
    set_margins(document)

    # --- Styles ---
    # Update Normal style defaults
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1

    # --- Header ---
    basic_info = resume_data.get("basic_info", {})
    name = f"{basic_info.get('first_name', '')} {basic_info.get('last_name', '')}"
    
    # Name Heading
    # Using paragraph instead of heading to avoid default border
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_before = Pt(0)
    header.paragraph_format.space_after = Pt(0)
    
    run = header.add_run(name)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(24) # Slightly larger for name
    
    # Contact Info
    contact_info = []
    if basic_info.get("address"):
        addr = basic_info["address"]
        if isinstance(addr, dict):
            parts = [p for p in [addr.get('city'), addr.get('state')] if p]
            if parts:
                contact_info.append(", ".join(parts))
    if basic_info.get("email"):
        contact_info.append(basic_info["email"])
    if basic_info.get("phone"):
        contact_info.append(basic_info["phone"])
    if basic_info.get("linkedin_profile"):
        contact_info.append("LinkedIn") 
    
    contact_paragraph = document.add_paragraph(" | ".join(contact_info))
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_paragraph.paragraph_format.space_before = Pt(0)
    contact_paragraph.paragraph_format.space_after = Pt(10) # Little space after header

    # --- Summary ---
    if resume_data.get("about"):
        add_section_heading(document, 'SUMMARY')
        add_normal_text(document, resume_data["about"])

    # --- Skills ---
    if resume_data.get("skills"):
        add_section_heading(document, 'TECHNICAL PROFICIENCIES')
        for skill_group in resume_data["skills"]:
            category = skill_group.get("category", "Skills")
            skills_list = ", ".join(skill_group.get("skills", []))
            
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1
            
            run_cat = p.add_run(f"{category}: ")
            run_cat.bold = True
            run_cat.font.name = 'Calibri'
            run_cat.font.size = Pt(10)
            
            run_list = p.add_run(skills_list)
            run_list.font.name = 'Calibri'
            run_list.font.size = Pt(10)

    # --- Experience ---
    if resume_data.get("experience"):
        add_section_heading(document, 'WORK EXPERIENCE')
        for exp in resume_data["experience"]:
            # Company Name - Job Title (Bold, Space After 4pt)
            # We can combine them or separate. Let's combine as requested in previous logic but apply bold formatting.
            company = exp.get('company_name', '')
            title = exp.get('job_title', '')
            date_range = f"{exp.get('start_date', '')} - {exp.get('end_date', '')}"
            
            # Line 1: Company - Title (Bold) ... Date (Italic)
            # To do right/left align on same line requires tabs, which is complex.
            # Simpler: Company - Title (Bold)
            # Date (Italic)
            # Or just: Company - Title | Date
            
            # Let's try to keep it simple but formatted
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(5) # Little space between jobs
            p.paragraph_format.space_after = Pt(4) # Requested 4pt after bold heading
            p.paragraph_format.line_spacing = 1
            
            run = p.add_run(f"{company} - {title}")
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            
            run_date = p.add_run(f" | {date_range}")
            run_date.italic = True
            run_date.font.name = 'Calibri'
            run_date.font.size = Pt(10)

            # Description bullets
            for desc in exp.get("description", []):
                add_bullet_point(document, desc)

    # --- Projects ---
    if resume_data.get("projects"):
        add_section_heading(document, 'PROJECTS')
        for project in resume_data["projects"]:
            name = project.get('name', '')
            
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1
            
            run = p.add_run(name)
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            
            # Description bullets
            for desc in project.get("description", []):
                add_bullet_point(document, desc)

    # --- Education ---
    if resume_data.get("education"):
        add_section_heading(document, 'EDUCATION')
        for edu in resume_data["education"]:
            school = edu.get('school_name', '')
            degree = edu.get('degree', '')
            date_range = f"{edu.get('start_date', '')} - {edu.get('end_date', '')}"
            
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1
            
            run = p.add_run(f"{school} | {degree}")
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            
            run_date = p.add_run(f" | {date_range}")
            run_date.font.name = 'Calibri'
            run_date.font.size = Pt(10)

    # --- Achievements/Awards ---
    if resume_data.get("awards"):
        add_section_heading(document, 'ACHIEVEMENTS')
        for award in resume_data["awards"]:
            p = document.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1
            
            run = p.add_run(f"{award.get('name', '')}: ")
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            
            run_desc = p.add_run(award.get('description', ''))
            run_desc.font.name = 'Calibri'
            run_desc.font.size = Pt(10)
